import os
import docx
import re

# Function to remove punctuation marks from text
def remove_punctuation(text):
    # Define the pattern to match punctuation marks
    punctuation_pattern = r'[^\w\s]'
    
    # Remove punctuation marks using regex
    text_without_punctuation = re.sub(punctuation_pattern, '', text)
    
    return text_without_punctuation

# Function to process .docx files in a directory and save processed files into a new folder
def process_files(input_folder, output_folder):
    # Create the output folder if it doesn't exist
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Process each .docx file in the input folder
    for filename in os.listdir(input_folder):
        if filename.endswith('.docx') and not filename.startswith('~$'):  # Ignore temporary files
            input_path = os.path.join(input_folder, filename)
            output_path = os.path.join(output_folder, filename)

            # Open the .docx file
            try:
                doc = docx.Document(input_path)

                # Extract text from the document
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"

                # Remove punctuation marks
                processed_text = remove_punctuation(text)

                # Create a new document
                new_doc = docx.Document()
                new_doc.add_paragraph(processed_text)

                # Save the new document
                new_doc.save(output_path)
            except Exception as e:
                print(f"Error processing file '{filename}': {e}")


# Specify input and output folders
input_folder = 'CFashionBeauty'  # Relative path from the current working directory
output_folder = 'ProcessedCFashionBeauty'  # Relative path from the current working directory

# Get the absolute paths
current_directory = os.path.dirname(os.path.abspath(__file__))
input_folder = os.path.join(current_directory, input_folder)
output_folder = os.path.join(current_directory, output_folder)

# Process files
process_files(input_folder, output_folder)
